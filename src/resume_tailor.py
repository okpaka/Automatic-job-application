from docx import Document
import spacy
import re
from nltk.corpus import stopwords
import nltk
import yaml

nltk.download('stopwords')

class ResumeTailor:
    def __init__(self):
        self.nlp = spacy.load("en_core_web_sm")
        self.stop_words = set(stopwords.words('english'))
        
        with open('config/config.yaml') as f:
            self.config = yaml.safe_load(f)
    
    def tailor_resume(self, job_description, output_path):
        """Customize resume based on job description"""
        # Load template resume
        doc = Document(self.config['resume_template_path'])
        
        # Extract keywords from job description
        keywords = self._extract_keywords(job_description)
        
        # Modify resume sections based on keywords
        self._modify_resume(doc, keywords)
        
        # Save tailored resume
        doc.save(output_path)
        return output_path
    
    def _extract_keywords(self, text):
        """Extract important keywords from job description"""
        doc = self.nlp(text.lower())
        
        # Get nouns and proper nouns that aren't stop words
        keywords = [
            token.text for token in doc 
            if not token.is_stop 
            and not token.is_punct 
            and (token.pos_ in ['NOUN', 'PROPN'] or token.text in ['python', 'java', 'sql'])
        ]
        
        # Count keyword frequency
        keyword_counts = {}
        for word in keywords:
            if word in keyword_counts:
                keyword_counts[word] += 1
            else:
                keyword_counts[word] = 1
        
        # Sort by frequency
        sorted_keywords = sorted(keyword_counts.items(), key=lambda x: x[1], reverse=True)
        return [kw[0] for kw in sorted_keywords[:20]]  # Return top 20 keywords
    
    def _modify_resume(self, doc, keywords):
        """Modify resume document to highlight keywords"""
        for paragraph in doc.paragraphs:
            text = paragraph.text.lower()
            
            # Highlight keywords in the text
            for keyword in keywords:
                if keyword in text:
                    # Find the keyword in the paragraph (case insensitive)
                    for match in re.finditer(re.escape(keyword), text, re.IGNORECASE):
                        start, end = match.start(), match.end()
                        # Clear the existing run
                        paragraph.clear()
                        # Add the text with keyword highlighted
                        paragraph.add_run(text[:start])
                        run = paragraph.add_run(text[start:end])
                        run.bold = True
                        run.underline = True
                        paragraph.add_run(text[end:])
                        break